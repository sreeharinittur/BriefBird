from flask import Flask, render_template, request
from transformers import BartTokenizer, BartForConditionalGeneration
from rouge import Rouge
from docx import Document
import re

app = Flask(__name__)

# Load BART model and tokenizer
model_name = "facebook/bart-large-cnn"
tokenizer = BartTokenizer.from_pretrained(model_name)
model = BartForConditionalGeneration.from_pretrained(model_name)

def count_words(text):
    return len(text.split())

def has_tables(docx_file):
    doc = Document(docx_file)
    return any(table for table in doc.tables)

def has_code(text):
    # Define regular expressions to identify code patterns
    code_patterns = [
        r'```[\s\S]*?```',  # Code blocks enclosed in triple backticks (```)
        r'(?s)<code>.*?</code>',  # Code blocks enclosed in HTML <code> tags
        r'(?i)(?<![\w\d_])if\s*\(.+?\)',  # If statements
        r'(?i)(?<![\w\d_])for\s*\(.+?\)',  # For loops
        r'(?i)(?<![\w\d_])while\s*\(.+?\)',  # While loops
        r'(?i)(?<![\w\d_])def\s+\w+\s*\(.+?\)',  # Function definitions
        r'(?i)(?<![\w\d_])class\s+\w+\s*:',  # Class definitions
        r'(?i)(?<![\w\d_])try\s*:',  # Try-except blocks
        r'(?i)(?<![\w\d_])except\s*:',  # Except blocks
        r'(?i)(?<![\w\d_])elif\s*\(.+?\)',  # Elif statements
        r'(?i)(?<![\w\d_])else\s*:',  # Else statements
        r'(?<![\w\d_])\s*=\s*.+?'  # Variable assignments
    ]
    
    # Check if any code pattern is found in the text
    for pattern in code_patterns:
        if re.search(pattern, text):
            return True
    return False

def summarize_text(text):
    input_ids = tokenizer.encode("summarize: " + text, return_tensors="pt", max_length=1024, truncation=True)
    summary_ids = model.generate(input_ids, max_length=150, min_length=50, length_penalty=2.0, num_beams=4, early_stopping=True)
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    return summary

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/contact")
def contact():
    return render_template("contact.html")

@app.route("/summarize", methods=["POST"])
def summarize():
    text = request.form["text"]
    selected_length = int(request.form["length"])

    # Check if the length of the input text is less than 30 words or it's too short
    if len(text.split()) < 30 or len(text) < 30:  # You can adjust the length threshold
        return render_template("index.html", text=text, summary=None, error="Please enter more meaningful text for a summary.", error_color="red")

    # Tokenize and encode the input text
    inputs = tokenizer.encode("summarize: " + text, return_tensors="pt", max_length=1024, truncation=True)

    # Generate the summary using BART
    summary_ids = model.generate(inputs, max_length=2 * selected_length, min_length=selected_length, length_penalty=2.0, num_beams=4, early_stopping=True)
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

    rouge = Rouge()
    rouge_scores = rouge.get_scores(summary, text)[0]

    return render_template("index.html", text=text, summary=summary, error=None, error_color="green",rouge_precision=rouge_scores['rouge-1']['p'], rouge_recall=rouge_scores['rouge-1']['r'], rouge_f1=rouge_scores['rouge-1']['f'])




@app.route('/doc', methods=['GET', 'POST'])
def doc():
    summary = None
    rouge_scores = None
    error_message = None

    if request.method == 'POST':
        if 'file' in request.files:
            docx_file = request.files['file']
            if docx_file and docx_file.filename.endswith('.docx'):
                text = extract_text_from_docx(docx_file)
                if has_tables(docx_file):
                    error_message = "Document contains tables and cannot be processed."
                elif has_code(text):
                    error_message = "Document contains code and cannot be processed."
                else:
                    if count_words(text) < 100:
                        error_message = "Document must contain at least 100 words."
                    elif count_words(text) > 1000:
                        error_message = "This document is too big to process!!!"
                    else:
                        summary = summarize_text(text)
                        rouge = Rouge()
                        rouge_scores = rouge.get_scores(summary, text)[0]
    if error_message:
        return render_template('doc.html', summary=None, rouge_scores=None, error_message=error_message)

    return render_template('doc.html', summary=summary, rouge_scores=rouge_scores, error_message=error_message)
if __name__ == "__main__":
    app.run(port=8000, debug=True)
